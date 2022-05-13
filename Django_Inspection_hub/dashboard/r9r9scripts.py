from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Cm, Inches
import os
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from datetime import date
import openpyxl
import os

def fromr8r9(report):

    # R8R9 Automation
    establishment_name = ''
    structure_type = ''
    height = ''
    ref_no = ''
    GA_photo = ''
    good_condition = ''
    location_of_structure = ''
    structure_defects_safe = ''
    structure_defects_unsafe = ''
    safe_access_equipment_defects = ''
    unsafe_access_equipment_defects = ''
    fitted_access = ''
    expiry_date = ''
    inspector_name = ''
    qualification = ''
    date_inspected = ''
    nearest_building = ''
    map_reference = ''
    use_description = ''
    owner = ''
    no_of_structures = ''
    inspection_frequency = ''
    climbing_facilities = ''
    access = ''
    note = ''
    previous_inspections = {}
    urgent_repairs = {}
    routine_repairs = {}
    preventative_maintenance_repairs = {}
    repair_counter = 0

    document = Document(report)

    for table in document.tables:
        row_count = len(table.rows)
        col_count = len(table.columns)
        for i in range(row_count):
            cells = table.rows[i].cells
            for cell in cells:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    for table in document.tables:
        if table.cell(0, 0).text == '1' or table.cell(1, 0).text == '1':
            row_count = len(table.rows)
            col_count = len(table.columns)
            for i in range(row_count):
                for j in range(col_count):
                    if table.rows[i].cells[j - 1].text == 'Establishment name' and table.rows[i].cells[j] \
                            .text != 'Establishment name':
                        establishment_name = table.rows[i].cells[j].text
                        continue
                    if table.rows[i].cells[j - 1].text == 'Type of structure' and table.rows[i].cells[j] \
                            .text != 'Type of structure':
                        structure_type = table.rows[i].cells[j].text
                        continue
                    if table.rows[i].cells[j - 1].text == 'Height in metres' and table.rows[i].cells[j] \
                            .text != 'Height in metres':
                        height = table.rows[i].cells[j].text
                        continue
                    if table.rows[i].cells[j - 1].text == 'Structure reference number' and table.rows[i].cells[j] \
                            .text != 'Structure reference number':
                        ref_no = table.rows[i].cells[j].text
                        continue
                    if table.rows[i].cells[j - 1].text == 'Location of structure' and table.rows[i].cells[j] \
                            .text != 'Location of structure':
                        location_of_structure = table.rows[i].cells[j].text
                        continue
                    if 'Defects which do NOT make the structure UNSAFE.' in table.rows[i].cells[
                        j - 1].text and 'Defects ' \
                                        'which do NOT make the structure UNSAFE.' not in table.rows[i].cells[j].text:
                        structure_defects_safe = table.rows[i].cells[j].text
                        continue
                    if 'Defects which make the structure UNSAFE' in table.rows[i].cells[j - 1].text and 'Defects which ' \
                                                                                                        'make the structure UNSAFE' not in \
                            table.rows[i].cells[j].text:
                        structure_defects_unsafe = table.rows[i].cells[j].text
                        continue
                    if '(Y/N)' in table.rows[i].cells[j - 1].text and '(Y/N)' not in table.rows[i].cells[j].text:
                        fitted_access = table.rows[i].cells[j].text
                        safe_access_equipment_defects = table.rows[i + 1].cells[j].text
                        unsafe_access_equipment_defects = table.rows[i + 2].cells[j].text
                        continue
                    if 'This certificate remains valid until' in table.rows[i].cells[j - 1].text and 'This certificate ' \
                                                                                                     'remains valid until' not in \
                            table.rows[i].cells[j].text:
                        expiry_date = table.rows[i].cells[j].text
                        continue
                    if 'Name of Competent Person' in table.rows[i].cells[j - 1].text and 'Name of ' \
                                                                                         'Competent Person' not in \
                            table.rows[i].cells[j].text:
                        inspector_name = table.rows[i].cells[j].text
                        qualification = table.rows[i + 1].cells[j].text
                        date_inspected = table.rows[i + 4].cells[j].text
                        continue
        elif 'Structure Location' in table.cell(0, 0).text:
            row_count = len(table.rows)
            col_count = len(table.columns)
            for i in range(row_count):
                for j in range(col_count):
                    if 'Nearest Building' in table.rows[i].cells[j].text:
                        nearest_building = table.rows[i].cells[j + 1].text
                        continue
                    if 'O.S' in table.rows[i].cells[j].text:
                        map_reference = table.rows[i].cells[j + 1].text
                        continue
                    if 'Description of Use' in table.rows[i].cells[j].text:
                        use_description = table.rows[i].cells[j + 1].text
                        continue
                    if 'Owner/Client' in table.rows[i].cells[j].text:
                        owner = table.rows[i].cells[j + 1].text
                        continue
                    if 'No. of Structures' in table.rows[i].cells[j].text:
                        no_of_structures = table.rows[i].cells[j + 1].text
                        continue
                    if 'Inspection Frequency' in table.rows[i].cells[j].text:
                        inspection_frequency = table.rows[i].cells[j + 1].text
                        continue
                    if 'Climbing Facilities' in table.rows[i].cells[j].text:
                        climbing_facilities = table.rows[i].cells[j + 1].text
                        continue
                    if 'Access to Structure' in table.rows[i].cells[j].text:
                        access = table.rows[i].cells[j + 1].text
                        continue
                    if 'Note' in table.rows[i].cells[j].text:
                        note = table.rows[i].cells[j + 1].text
                        continue
        elif 'Date Inspected' in table.cell(0, 0).text:
            row_count = len(table.rows)
            col_count = len(table.columns)
            for i in range(1, row_count):
                previous_inspections[f'inspection{i}'] = {'Date Insepcted': table.cell(i, 0).text,
                                                          'Round No.': table.cell(i, 1).text,
                                                          'Date of R8': table.cell(i, 2).text,
                                                          'Structure Unsafe': table.cell(i, 3).text,
                                                          'Good Condition': table.cell(i, 4).text,
                                                          'Maintenance Required': table.cell(i, 5).text,
                                                          'Maintenance Prior': table.cell(i, 6).text,
                                                          'Next Inspection Date': table.cell(i, 7).text}
            break
        else:
            row_count = len(table.rows)
            col_count = len(table.columns)
            for i in range(row_count):
                for j in range(col_count):
                    if 'Element' in table.rows[i].cells[j - 1].text and 'Description ' \
                                                                        'of defect' in table.rows[i].cells[j].text and \
                                                                        repair_counter == 0:
                        row_counter = 1
                        end = ''
                        while end == '':
                            if 'Total (carried to summary)' in table.rows[i + row_counter].cells[j - 1].text:
                                end = 'End'
                                continue
                            if table.rows[i + row_counter].cells[0].text != '' and 'Total ' \
                                                                                   '(carried to summary)' not in \
                                    table.rows[i + row_counter].cells[0].text:
                                urgent_repairs[f'repair_{row_counter}'] = {
                                    'Element': table.rows[i + row_counter].cells[0].text,
                                    'Defect': table.rows[i + row_counter].cells[j].text,
                                    'Cost': table.rows[i + row_counter].cells[col_count - 1].text}
                            row_counter += 1
                        repair_counter += 1
                        continue
                    if 'Element' in table.rows[i].cells[j - 1].text and 'Description ' \
                                                                        'of defect' in table.rows[i].cells[
                        j].text and repair_counter == 1:
                        row_counter = 1
                        end = ''
                        while end == '':
                            if 'Total (carried to summary)' in table.rows[i + row_counter].cells[j - 1].text:
                                end = 'End'
                                continue
                            if table.rows[i + row_counter].cells[0].text != '' and 'Total ' \
                                                                                   '(carried to summary)' not in \
                                    table.rows[i + row_counter].cells[0].text:
                                routine_repairs[f'repair_{row_counter}'] = {
                                    'Element': table.rows[i + row_counter].cells[0].text,
                                    'Defect': table.rows[i + row_counter].cells[j].text,
                                    'Cost': table.rows[i + row_counter].cells[col_count - 1].text}
                            row_counter += 1
                        repair_counter += 1
                        continue
                    if 'Element' in table.rows[i].cells[j - 1].text and 'Description ' \
                                                                        'of defect' in table.rows[i].cells[
                        j].text and repair_counter == 2:
                        row_counter = 1
                        end = ''
                        while end == '':
                            if 'Total (carried to summary)' in table.rows[i + row_counter].cells[j - 1].text:
                                end = 'End'
                                continue
                            if table.rows[i + row_counter].cells[0].text != '' and 'Total ' \
                                                                                   '(carried to summary)' not in \
                                    table.rows[i + row_counter].cells[0].text:
                                preventative_maintenance_repairs[f'repair_{row_counter}'] = {
                                    'Element': table.rows[i + row_counter].cells[0].text,
                                    'Defect': table.rows[i + row_counter].cells[j].text,
                                    'Cost': table.rows[i + row_counter].cells[col_count - 1].text}
                            row_counter += 1
                        repair_counter += 1
                        continue

    for key in previous_inspections:
        good_condition = previous_inspections[key]['Good Condition']

    output = {'establishment_name': establishment_name,
              'structure_type': structure_type,
              'height': height,
              'ref_no': ref_no,
              'GA_photo': GA_photo,
              'good_condition': good_condition,
              'location_of_structure': location_of_structure,
              'structure_defects_safe': structure_defects_safe,
              'structure_defects_unsafe': structure_defects_unsafe,
              'safe_access_equipment_defects': safe_access_equipment_defects,
              'unsafe_access_equipment_defects':unsafe_access_equipment_defects,
              'fitted_access': fitted_access,
              'expiry_date': expiry_date,
              'inspector_name': inspector_name,
              'qualification': qualification,
              'date_inspected': date_inspected,
              'nearest_building': nearest_building,
              'map_reference': map_reference,
              'use_description': use_description,
              'owner': owner,
              'no_of_structures': no_of_structures,
              'inspection_frequency': inspection_frequency,
              'climbing_facilities': climbing_facilities,
              'access': access,
              'note': note,
              'previous_inspections': previous_inspections,
              'urgent_repairs': urgent_repairs,
              'routine_repairs': routine_repairs,
              'preventative_maintenance_repairs': preventative_maintenance_repairs}

    return output


def getpreviousinspectiondata(report):

    previous_inspections = {}

    document = Document(report)

    for table in document.tables:
        if 'Date Inspected' in table.cell(0, 0).text:
            row_count = len(table.rows)
            col_count = len(table.columns)
            for i in range(1, row_count):
                previous_inspections[f'inspection{i}'] = {'Date Insepcted': table.cell(i, 0).text,
                                                          'Round No.': table.cell(i, 1).text,
                                                          'Date of R8': table.cell(i, 2).text,
                                                          'Structure Unsafe': table.cell(i, 3).text,
                                                          'Good Condition': table.cell(i, 4).text,
                                                          'Maintenance Required': table.cell(i, 5).text,
                                                          'Maintenance Prior': table.cell(i, 6).text,
                                                          'Next Inspection Date': table.cell(i, 7).text}

    return previous_inspections
