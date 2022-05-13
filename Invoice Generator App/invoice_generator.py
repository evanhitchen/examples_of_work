from tkinter import Tk,Label,Button,filedialog, Frame
from pdf2docx import Converter
from docx2pdf import convert
from docx import *
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pdfplumber
import sys


def main():
    # Raise frame function
    def raise_frame(frame):
        frame.tkraise()

    def exit_script():
        sys.exit()

    global pdf_file

    def popup_bonus():
        popup = Tk()
        popup.wm_title("Warning")
        popup.geometry("200x100")
        label = Label(popup, text='Please select invoice PDF')
        label.pack(side="top", fill="x", pady=10)
        B1 = Button(popup, text="Okay", command = popup.destroy)
        B1.pack()
        popup.mainloop()

    def invoice():
        invoice_generator(pdf_file)

    def invoice_generator(pdf_file):

        if pdf_file == '':
            print()
            popup_bonus()
        else:
            document_names = []
            pdf = pdfplumber.open(pdf_file)
            for page_no in range(len(pdf.pages)):
                # try:
                page = pdf.pages[page_no]
                text = page.extract_text()
                text_list = text.split('\n')
                pdf.close()

                if 'INVOICE DRAFT' in text_list:
                    start = text_list.index('INVOICE DRAFT')
                    template_start = 'draft_invoice'
                else:
                    start = text_list.index('INVOICE')
                    template_start = 'invoice'
                for i in text_list:
                    if 'VAT Reg Number:' in i:
                        index_item = i
                        break
                end = text_list.index(index_item)
                address_section = text_list[start+1:end]
                try:
                    address_section.remove('WITHOUT PREJUDICE')
                    template_doc = f'Ref/{template_start}_template_without_prejudice.docx'
                except:
                    template_doc = f'Ref/{template_start}_template.docx'

                address = ''
                for i in address_section:
                    address = address + i + '\n'

                for i in text_list:
                    if 'VAT Reg Number:' in i:
                        vat_reg_number = i.split('VAT Reg Number: ')[1]
                    elif 'Our Ref:' in i:
                        our_ref_section = i.split(' Tax Point Date: ')
                        tax_point_date = i.split(' Tax Point Date: ')[1]
                        our_ref = our_ref_section[0].split('Our Ref: ')[1]
                    elif 'Invoice Number:' in i:
                        invoice_no = i.split('Invoice Number:')[1]
                        if len(invoice_no) > 1:
                            if invoice_no[0] == ' ':
                                invoice_no = invoice_no[1:]
                    elif 'Tenant:' in i:
                        tenant = i.split('Tenant: ')[1]
                    elif 'Tenant Ref:' in i:
                        tenant_ref = i.split('Tenant Ref: ')[1]
                    elif 'Lease Ref:' in i:
                        lease_ref = i.split('Lease Ref: ')[1]
                    elif 'Net ' == i[0:4]:
                        total_fee = i.split(' ')[1]
                    elif 'Current Balance' in i:
                        current_balance_text = i.split('Current Balance: ')[1]
                        current_balance = current_balance_text.split(' RH16 1PG')[0]

                cv = Converter(pdf_file)
                tables = cv.extract_tables(start=page_no, end=(page_no+1))
                cv.close()
                table_rows = []
                row_count = len(tables[0])-1
                for i in range(row_count):
                    table_rows.append(tables[0][i+1])

                document = Document(template_doc)

                for table in document.tables:
                    if table.rows[0].cells[0].text == 'Property':
                        # table.rows[1].cells[1].text = 'hi'
                        for i in range(7):
                            table.rows[1].cells[i].text = table_rows[0][i]
                            paragraphs = table.rows[1].cells[i].paragraphs
                            for paragraph in paragraphs:
                                paragraph.style = document.styles['Table Paragraph']
                                if i == 5 or i == 6 or i == 4:
                                    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        if len(table_rows) > 1:
                            for j in range(len(table_rows)-1):
                                new_row = table.add_row()
                                for i in range(7):
                                    new_row.cells[i].text = table_rows[2+j][i]
                                    paragraphs = new_row.cells[i].paragraphs
                                    for paragraph in paragraphs:
                                        paragraph.style = document.styles['Table Paragraph']
                                        if i == 5 or i == 6 or i == 4:
                                            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                if 'Address_input' in paragraph.text:
                                    paragraph.text = address
                                if 'Our_ref' in paragraph.text:
                                    paragraph.text = our_ref
                                if 'Vat_reg' in paragraph.text:
                                    paragraph.text = vat_reg_number
                                if 'Tax_point' in paragraph.text:
                                    paragraph.text = tax_point_date
                                if 'Invoice_no' in paragraph.text:
                                    paragraph.text = invoice_no
                                if 'Tenant_ref' in paragraph.text:
                                    paragraph.text = tenant_ref
                                if 'Tenant_name' in paragraph.text:
                                    paragraph.text = tenant
                                if 'Lease_ref' in paragraph.text:
                                    paragraph.text = lease_ref
                                if 'Invoice_no' in paragraph.text:
                                    paragraph.text = invoice_no
                                if 'balance'in paragraph.text:
                                    paragraph.text = current_balance
                                if 'Net_value' in paragraph.text:
                                    paragraph.text = total_fee

                # Save document
                # Save document
                doc_name = f'Invoice Number {our_ref}.docx'
                counter = 1
                while doc_name in document_names:
                    doc_name = f'Invoice Number {our_ref}({counter}).docx'
                    counter += 1
                document_names.append(doc_name)
                document.save(doc_name)
                convert(doc_name)
                # Change label contents
                label_file_explorer.configure(
                    text="File Opened: " + pdf_file + '\n' + f'Page {page_no + 1} of '
                                                             f'{len(pdf.pages)} converted')

    # Function for opening the
    # file explorer window
    def browsefiles():
        filename = filedialog.askopenfilename(initialdir="/",
                                              title="Select a File",
                                              filetypes=(("PDF Files",
                                                          "*.pdf*"),
                                                         ("all files",
                                                          "*.*")))

        # Change label contents
        label_file_explorer.configure(text="File Opened: " + filename)

        global pdf_file
        pdf_file = filename

    # Create the root window
    window = Tk()
    frame1 = Frame(window)

    # Set window title
    window.title('File Explorer')

    # Set window size
    window.geometry("700x300")

    # Set window background color
    frame1.config(background="white")

    frame1.grid(row=0, column=0, sticky='news')

    # Create a File Explorer label
    label_file_explorer = Label(frame1,
                                text="File Explorer using Tkinter",
                                width=100, height=4,
                                fg="blue")

    button_explore = Button(frame1,
                            text="Browse Files",
                            command=browsefiles)

    invoice_generation = Button(frame1,
                            text="Generate Invoice(s)",
                            command=invoice)

    button_exit = Button(frame1,
                         text="Exit",
                         command=exit_script)

    # Grid method is chosen for placing
    # the widgets at respective positions
    # in a table like structure by
    # specifying rows and columns
    label_file_explorer.grid(column=1, row=5)

    button_explore.grid(column=1, row=2)

    invoice_generation.grid(column=1, row=3)

    button_exit.grid(column=1, row=4)

    raise_frame(frame1)
    # Let the window wait for any events
    window.mainloop()


if __name__ == '__main__':
    main()